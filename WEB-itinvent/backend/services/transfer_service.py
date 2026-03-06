"""
Transfer service: act generation, in-memory act registry and email sending.
"""
from __future__ import annotations

import asyncio
import logging
import mimetypes
import os
import re
import smtplib
import tempfile
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import Any, Optional
from uuid import uuid4

from dotenv import dotenv_values

logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_TEMPLATE_PATH = PROJECT_ROOT / "templates" / "docx_transfer_act.docx"
DEFAULT_ACTS_DIR = PROJECT_ROOT / os.getenv("TRANSFER_ACTS_DIR", "transfer_acts")
BOT_ENV_PATH = PROJECT_ROOT / ".env"
BOT_ENV_VALUES = dotenv_values(str(BOT_ENV_PATH)) if BOT_ENV_PATH.exists() else {}

_ACT_STORE: dict[str, dict[str, Any]] = {}


def _read_env(name: str, default: Optional[str] = None) -> Optional[str]:
    value = os.getenv(name)
    if value is not None and str(value).strip() != "":
        return str(value).strip()
    fallback = BOT_ENV_VALUES.get(name)
    if fallback is not None and str(fallback).strip() != "":
        return str(fallback).strip()
    return default


def _is_private_smtp_host(host: str) -> bool:
    h = (host or "").strip().lower()
    return (
        h.startswith("10.")
        or h.startswith("192.168.")
        or h.startswith("172.")
        or h in {"localhost", "127.0.0.1"}
    )


def _is_placeholder(value: str) -> bool:
    v = (value or "").strip().lower()
    return v in {"your_email@company.com", "your_email_password", "changeme", "change-me"}


def _to_bool(value: Optional[str], default: bool = False) -> bool:
    if value is None:
        return default
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


class _EmailSender:
    """Minimal SMTP sender for transfer acts."""

    def __init__(self) -> None:
        self.smtp_server = _read_env("SMTP_SERVER", "localhost") or "localhost"
        self.smtp_port = int(_read_env("SMTP_PORT", "25") or "25")
        self.email = (
            _read_env("EMAIL_ADDRESS")
            or _read_env("SMTP_FROM_EMAIL")
            or "noreply@localhost"
        )
        username = _read_env("SMTP_USERNAME")
        password = _read_env("EMAIL_PASSWORD") or _read_env("SMTP_PASSWORD") or ""

        self.username = username or self.email
        self.password = "" if _is_placeholder(password) else password

        use_auth_env = _read_env("SMTP_USE_AUTH")
        if use_auth_env is not None:
            self.use_auth = _to_bool(use_auth_env, default=False)
        else:
            # Behaves like bot/email_sender.py: local/private SMTP -> no AUTH
            self.use_auth = not _is_private_smtp_host(self.smtp_server)

        # AUTH without password is invalid
        if self.use_auth and not self.password:
            self.use_auth = False

        self.use_tls = _to_bool(_read_env("SMTP_USE_TLS"), default=False) and self.use_auth

    def send_files(
        self,
        recipient_email: str,
        files: dict[str, str],
        subject: str,
        body: str,
    ) -> bool:
        msg = MIMEMultipart()
        msg["From"] = self.email
        msg["To"] = recipient_email
        msg["Subject"] = subject
        msg.attach(MIMEText(body or "", "plain", "utf-8"))

        attached = 0
        for _, path in files.items():
            if not path or not os.path.exists(path):
                continue
            filename = os.path.basename(path)
            mime_type, _ = mimetypes.guess_type(path)
            if not mime_type:
                mime_type = "application/octet-stream"
            main, sub = mime_type.split("/", 1)

            with open(path, "rb") as stream:
                part = MIMEBase(main, sub)
                part.set_payload(stream.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
            msg.attach(part)
            attached += 1

        if attached == 0:
            return False

        try:
            with smtplib.SMTP(self.smtp_server, self.smtp_port, timeout=30) as server:
                if self.use_tls:
                    server.ehlo()
                    server.starttls()
                    server.ehlo()
                if self.use_auth and self.password:
                    server.login(self.username, self.password)
                server.sendmail(self.email, [recipient_email], msg.as_string())
            return True
        except smtplib.SMTPNotSupportedError as exc:
            # Fallback for local SMTP servers without AUTH extension
            logger.warning("SMTP AUTH not supported, retrying without AUTH: %s", exc)
            try:
                with smtplib.SMTP(self.smtp_server, self.smtp_port, timeout=30) as server:
                    server.sendmail(self.email, [recipient_email], msg.as_string())
                return True
            except Exception as retry_exc:
                logger.error("Email send retry failed: %s", retry_exc)
                return False
        except Exception as exc:
            logger.error("Email send failed: %s", exc)
            return False


def _safe_filename(name: str) -> str:
    value = (name or "Unknown").strip()
    value = re.sub(r"\s+", "_", value)
    value = re.sub(r"[^\w.-]", "", value)
    value = value.strip("._")
    return value[:48] or "Unknown"


def _to_inv_text(value: Any) -> str:
    if value is None:
        return ""
    try:
        return str(int(round(float(value))))
    except (TypeError, ValueError):
        return str(value)


def _build_docx_act(
    old_employee: str,
    new_employee: str,
    new_employee_dept: str,
    items: list[dict[str, Any]],
) -> tuple[Path, str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx not installed") from exc

    template_path = DEFAULT_TEMPLATE_PATH
    if not template_path.exists():
        raise RuntimeError(f"Act template not found: {template_path}")

    DEFAULT_ACTS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))
    now = datetime.now()
    date_text = now.strftime("%d.%m.%Y")

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "{{DATE}}" in text:
            paragraph.text = text.replace("{{DATE}}", date_text)
            text = paragraph.text
        if "{{TO_EMPLOYEE}}" in text:
            paragraph.text = text.replace("{{TO_EMPLOYEE}}", str(new_employee))
            text = paragraph.text
        if "{{FROM_EMPLOYEE}}" in text:
            paragraph.text = text.replace("{{FROM_EMPLOYEE}}", str(old_employee))

    if doc.tables:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        table = doc.tables[0]
        if len(table.rows) > 1:
            template_row = table.rows[1]
            table._element.remove(template_row._element)

        for idx, item in enumerate(items, 1):
            row = table.add_row()
            cells_data = [
                str(idx),
                str(item.get("type_name") or ""),
                str(item.get("model_name") or ""),
                str(item.get("serial_no") or ""),
                str(item.get("part_no") or ""),
                str(new_employee_dept or ""),
                _to_inv_text(item.get("inv_no")),
            ]
            for cell_idx, cell_text in enumerate(cells_data):
                cell = row.cells[cell_idx]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tc = cell._element
                tc_pr = tc.get_or_add_tcPr()
                valign = OxmlElement("w:vAlign")
                valign.set(qn("w:val"), "center")
                tc_pr.append(valign)

    timestamp = now.strftime("%Y%m%d_%H%M%S")
    safe_old = _safe_filename(old_employee)
    docx_name = f"transfer_act_{timestamp}_{safe_old}.docx"
    docx_path = DEFAULT_ACTS_DIR / docx_name
    doc.save(str(docx_path))
    return docx_path, "docx"


def _try_convert_to_pdf(docx_path: Path) -> tuple[Path, str]:
    try:
        from docx2pdf import convert  # type: ignore
    except ImportError:
        return docx_path, "docx"
    except Exception:
        return docx_path, "docx"

    pdf_path = docx_path.with_suffix(".pdf")
    try:
        convert(str(docx_path), str(pdf_path))
    except Exception as exc:
        logger.warning("DOCX->PDF conversion failed: %s", exc)
        return docx_path, "docx"

    if pdf_path.exists():
        try:
            docx_path.unlink(missing_ok=True)
        except Exception:
            pass
        return pdf_path, "pdf"

    return docx_path, "docx"


def get_act_record(act_id: str) -> Optional[dict[str, Any]]:
    """Get act record by ID."""
    return _ACT_STORE.get(str(act_id or "").strip())


def generate_transfer_acts(
    transferred_items: list[dict[str, Any]],
    new_employee_name: str,
    new_employee_dept: str,
    new_employee_email: Optional[str],
    db_id: Optional[str],
) -> list[dict[str, Any]]:
    """
    Generate transfer acts grouped by old employee.

    Returns list with API-ready act metadata.
    """
    grouped: dict[str, list[dict[str, Any]]] = {}
    for item in transferred_items:
        old_employee = (item.get("old_employee_name") or "Без владельца").strip() or "Без владельца"
        grouped.setdefault(old_employee, []).append(item)

    acts: list[dict[str, Any]] = []

    for old_employee, items in grouped.items():
        old_employee_email = None
        for item in items:
            candidate = (item.get("old_employee_email") or "").strip()
            if candidate:
                old_employee_email = candidate
                break

        try:
            file_path, file_type = _build_docx_act(
                old_employee=old_employee,
                new_employee=new_employee_name,
                new_employee_dept=new_employee_dept or "",
                items=items,
            )
            file_path, file_type = _try_convert_to_pdf(file_path)
        except Exception as exc:
            logger.error("Failed to generate act for %s: %s", old_employee, exc)
            continue

        act_id = str(uuid4())
        record = {
            "act_id": act_id,
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_type": file_type,
            "old_employee_name": old_employee,
            "old_employee_email": old_employee_email,
            "new_employee_name": new_employee_name,
            "new_employee_email": (new_employee_email or "").strip() or None,
            "equipment_count": len(items),
            "db_id": db_id,
            "created_at": datetime.now().isoformat(),
        }
        _ACT_STORE[act_id] = record
        acts.append(
            {
                "act_id": act_id,
                "old_employee": old_employee,
                "equipment_count": len(items),
                "file_name": file_path.name,
                "file_type": file_type,
            }
        )

    return acts


async def _send_files(
    recipient_email: str,
    records: list[dict[str, Any]],
    subject: str,
    body: str,
) -> bool:
    files: dict[str, str] = {}
    for idx, record in enumerate(records):
        file_path = record.get("file_path")
        if file_path and os.path.exists(file_path):
            files[f"act_{idx+1}"] = file_path
    if not files:
        return False

    sender = _EmailSender()
    return await asyncio.to_thread(
        sender.send_files,
        recipient_email=recipient_email,
        files=files,
        subject=subject,
        body=body,
    )


async def send_binary_file_email(
    *,
    recipient_email: str,
    file_name: str,
    file_bytes: bytes,
    subject: str,
    body: str,
) -> bool:
    """
    Send one in-memory file as email attachment.
    """
    safe_recipient = str(recipient_email or "").strip()
    if not safe_recipient or not file_bytes:
        return False

    safe_name = str(file_name or "").strip() or "act.pdf"
    safe_name = re.sub(r"[\\/:*?\"<>|]+", "_", safe_name).strip() or "act.pdf"
    temp_path = ""
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            temp_path = os.path.join(tmp_dir, safe_name)
            with open(temp_path, "wb") as tmp:
                tmp.write(bytes(file_bytes))

            sender = _EmailSender()
            return await asyncio.to_thread(
                sender.send_files,
                recipient_email=safe_recipient,
                files={"uploaded_act": temp_path},
                subject=str(subject or "").strip(),
                body=str(body or "").strip(),
            )
    except Exception:
        return False


async def send_transfer_acts_email(
    act_ids: list[str],
    mode: str,
    manual_email: Optional[str] = None,
    employee_email: Optional[str] = None,
) -> dict[str, Any]:
    """
    Send generated acts by email.

    mode:
      - old: send each act to old employee email
      - new: send all acts to new employee email
      - manual: send all acts to manual_email
      - employee: send all acts to employee_email
    """
    records: list[dict[str, Any]] = []
    errors: list[str] = []

    for act_id in act_ids:
        record = get_act_record(act_id)
        if not record:
            errors.append(f"Act not found: {act_id}")
            continue
        if not os.path.exists(record.get("file_path", "")):
            errors.append(f"Act file not found: {record.get('file_name', act_id)}")
            continue
        records.append(record)

    if not records:
        return {
            "success_count": 0,
            "failed_count": max(1, len(errors)),
            "errors": errors or ["No valid act files selected"],
        }

    success_count = 0
    failed_count = 0

    if mode == "old":
        recipient_map: dict[str, list[dict[str, Any]]] = {}
        for record in records:
            recipient = (record.get("old_employee_email") or "").strip()
            if not recipient:
                failed_count += 1
                errors.append(f"No email for old employee: {record.get('old_employee_name')}")
                continue
            recipient_map.setdefault(recipient, []).append(record)

        for recipient, recipient_records in recipient_map.items():
            subject = f"Акты приема-передачи ({datetime.now().strftime('%d.%m.%Y')})"
            body = "Во вложении акт приема-передачи оборудования."
            sent = await _send_files(recipient, recipient_records, subject, body)
            if sent:
                success_count += 1
            else:
                failed_count += 1
                errors.append(f"Failed to send to {recipient}")

    else:
        if mode == "new":
            recipient_email = (records[0].get("new_employee_email") or "").strip()
        elif mode == "manual":
            recipient_email = (manual_email or "").strip()
        elif mode == "employee":
            recipient_email = (employee_email or "").strip()
        else:
            recipient_email = ""

        if not recipient_email:
            return {
                "success_count": 0,
                "failed_count": 1,
                "errors": ["Recipient email is not set"],
            }

        subject = f"Акты приема-передачи ({datetime.now().strftime('%d.%m.%Y')})"
        body = "Во вложении акты приема-передачи оборудования."
        sent = await _send_files(recipient_email, records, subject, body)
        if sent:
            success_count += 1
        else:
            failed_count += 1
            errors.append(f"Failed to send to {recipient_email}")

    return {
        "success_count": success_count,
        "failed_count": failed_count,
        "errors": errors,
    }
