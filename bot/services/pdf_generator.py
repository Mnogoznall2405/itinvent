#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Сервис генерации PDF-актов приема-передачи оборудования
"""
import logging
import os
import asyncio
import time
import gc
from datetime import datetime
from typing import List, Dict, Any, Optional

logger = logging.getLogger(__name__)


def remove_file_with_retry(filepath: str, max_attempts: int = 5, delay: float = 0.5) -> bool:
    """
    Удаляет файл с механизмом повторных попыток при ошибке блокировки

    Параметры:
        filepath: Путь к удаляемому файлу
        max_attempts: Максимальное количество попыток
        delay: Задержка между попытками в секундах

    Возвращает:
        bool: True если файл удалён, False если все попытки неудачны
    """
    for attempt in range(max_attempts):
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                logger.info(f"Файл удалён: {filepath}")
                return True
            else:
                logger.debug(f"Файл уже不存在: {filepath}")
                return True
        except PermissionError as e:
            logger.warning(f"Попытка {attempt + 1}/{max_attempts}: Файл занят {filepath} - {e}")
            if attempt < max_attempts - 1:
                time.sleep(delay * (attempt + 1))  # Увеличиваем задержку с каждой попыткой
            else:
                logger.error(f"Не удалось удалить файл после {max_attempts} попыток: {filepath}")
                return False
        except Exception as e:
            logger.error(f"Ошибка удаления файла {filepath}: {e}")
            return False
    return False


def remove_word_temp_files(docx_path: str) -> None:
    """
    Удаляет временные файлы Word, если они есть

    Параметры:
        docx_path: Путь к DOCX файлу
    """
    try:
        directory, filename = os.path.split(docx_path)
        # Word создаёт временные файлы с префиксом ~$
        temp_file_pattern = f"~$ {filename}"
        temp_file_path = os.path.join(directory, temp_file_pattern)

        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
                logger.info(f"Временный файл Word удалён: {temp_file_path}")
            except Exception as e:
                logger.debug(f"Не удалось удалить временный файл Word: {e}")
    except Exception as e:
        logger.debug(f"Ошибка при поиске временных файлов Word: {e}")


async def generate_transfer_act_pdf(
    new_employee: str,
    new_employee_dept: str,
    old_employee: str,
    serials_data: List[Dict[str, Any]],
    db_name: str
) -> Optional[str]:
    """
    Генерирует PDF-акт приема-передачи оборудования из шаблона
    
    Параметры:
        new_employee: ФИО нового сотрудника
        new_employee_dept: Отдел нового сотрудника
        old_employee: ФИО старого сотрудника (от кого передается)
        serials_data: Список данных об оборудовании
        db_name: Название базы данных
        
    Возвращает:
        str: Путь к созданному PDF-файлу или None при ошибке
    """
    try:
        # Импортируем библиотеки
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            logger.error("Библиотека python-docx не установлена. Установите: pip install python-docx")
            return None
        
        try:
            from docx2pdf import convert
        except ImportError:
            logger.error("Библиотека docx2pdf не установлена. Установите: pip install docx2pdf")
            return None
        
        # Путь к шаблону
        template_path = "templates/docx_transfer_act.docx"
        if not os.path.exists(template_path):
            logger.error(f"Шаблон не найден: {template_path}")
            return None
        
        # Создаем директорию для актов если её нет
        acts_dir = 'transfer_acts'
        if not os.path.exists(acts_dir):
            os.makedirs(acts_dir)
        
        # Загружаем шаблон
        doc = Document(template_path)
        
        # Текущая дата
        current_date = datetime.now()
        date_str = current_date.strftime('%d.%m.%Y')
        
        # Заменяем плейсхолдеры в параграфах
        for paragraph in doc.paragraphs:
            if '{{DATE}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{DATE}}', date_str)
            if '{{TO_EMPLOYEE}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{TO_EMPLOYEE}}', str(new_employee))
            if '{{FROM_EMPLOYEE}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FROM_EMPLOYEE}}', old_employee)
        
        # Работаем с таблицей оборудования
        if doc.tables:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            table = doc.tables[0]
            
            # Удаляем строку-шаблон (вторую строку, индекс 1)
            if len(table.rows) > 1:
                template_row = table.rows[1]
                table._element.remove(template_row._element)
            
            # Добавляем строки с данными оборудования
            for idx, item in enumerate(serials_data, 1):
                equipment = item.get('equipment', {})
                serial = str(item.get('serial', 'Не указан'))
                
                # Получаем данные с обработкой null значений
                type_name = equipment.get('TYPE_NAME') or ''
                model_name = equipment.get('MODEL_NAME') or 'Не указано'
                # PART_NO - инвентарный номер, может быть null, поэтому проверяем и заменяем на пустую строку
                batch_no = equipment.get('PART_NO') or ''
                
                # Инвентарный номер - округляем до целого если это число
                inv_no = equipment.get('INV_NO')
                if inv_no is None or inv_no == '':
                    inv_no_str = ''
                else:
                    try:
                        # Пробуем преобразовать в число и округлить
                        inv_no_float = float(inv_no)
                        inv_no_str = str(int(round(inv_no_float)))
                    except (ValueError, TypeError):
                        inv_no_str = str(inv_no)
                
                # Добавляем строку
                row = table.add_row()
                
                # Заполняем ячейки
                # Используем отдел НОВОГО сотрудника (получателя), а не старого
                cells_data = [
                    str(idx),
                    str(type_name) if type_name else '',
                    str(model_name) if model_name else '',
                    serial,
                    str(batch_no),  # Теперь batch_no уже строка или пустая строка
                    str(new_employee_dept) if new_employee_dept else '',  # Отдел получателя
                    inv_no_str
                ]
                
                for cell_idx, cell_text in enumerate(cells_data):
                    cell = row.cells[cell_idx]
                    cell.text = cell_text
                    
                    # Выравнивание по центру и по середине
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Вертикальное выравнивание по центру
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), 'center')
                    tcPr.append(tcVAlign)
        
        # Импортируем sanitize_filename
        from bot.services.equipment_grouper import sanitize_filename
        
        # Сохраняем DOCX
        timestamp = current_date.strftime('%Y%m%d_%H%M%S')
        old_employee_sanitized = sanitize_filename(old_employee)
        docx_filename = f'transfer_act_{timestamp}_{old_employee_sanitized}.docx'
        docx_path = os.path.join(acts_dir, docx_filename)

        # Сохраняем DOCX
        doc.save(docx_path)
        logger.info(f"DOCX-акт создан: {docx_path}")

        # Освобождаем ресурсы перед конвертацией
        del doc
        gc.collect()

        # Конвертируем в PDF
        pdf_filename = f'transfer_act_{timestamp}_{old_employee_sanitized}.pdf'
        pdf_path = os.path.join(acts_dir, pdf_filename)

        try:
            # Запускаем конвертацию в отдельном потоке с таймаутом
            loop = asyncio.get_event_loop()

            # Конвертируем с таймаутом 60 секунд
            await asyncio.wait_for(
                loop.run_in_executor(None, convert, docx_path, pdf_path),
                timeout=60.0
            )

            logger.info(f"PDF-акт создан: {pdf_path}")

            # Даём время Word освободить файл после конвертации
            await asyncio.sleep(1.0)

            # Удаляем временный DOCX с механизмом повторных попыток
            remove_file_with_retry(docx_path, max_attempts=5, delay=0.5)

            # Также удаляем временные файлы Word (~$*)
            remove_word_temp_files(docx_path)

            return pdf_path

        except asyncio.TimeoutError:
            logger.error(f"Таймаут конвертации DOCX в PDF (превышено 60 секунд)")
            # Если конвертация не удалась по таймауту, возвращаем DOCX
            logger.warning(f"Возвращаем DOCX вместо PDF из-за таймаута: {docx_path}")
            return docx_path
        except Exception as e:
            logger.error(f"Ошибка конвертации DOCX в PDF: {e}")
            # Если конвертация не удалась, возвращаем DOCX
            logger.warning(f"Возвращаем DOCX вместо PDF: {docx_path}")
            return docx_path
        
    except Exception as e:
        logger.error(f"Ошибка генерации акта: {e}", exc_info=True)
        return None


async def generate_multiple_transfer_acts(
    new_employee: str,
    new_employee_dept: str,
    grouped_equipment: Dict[str, List[Dict[str, Any]]],
    db_name: str
) -> List[Dict[str, Any]]:
    """
    Генерирует отдельный PDF-акт для каждой группы оборудования
    
    Параметры:
        new_employee: ФИО нового сотрудника (получатель)
        new_employee_dept: Отдел нового сотрудника
        grouped_equipment: Словарь {old_employee: [equipment_list]}
        db_name: Название базы данных
    
    Возвращает:
        List[Dict]: Список информации о созданных актах:
            [
                {
                    "old_employee": "Иванов Иван Иванович",
                    "pdf_path": "transfer_acts/transfer_act_20240122_143000_Ivanov.pdf",
                    "filename": "transfer_act_20240122_143000_Ivanov.pdf",
                    "equipment_count": 2,
                    "success": True,
                    "error": None
                },
                ...
            ]
    """
    logger.info(f"Начало генерации множественных актов: {len(grouped_equipment)} групп")
    start_time = time.time()
    
    # Создаем задачи для параллельной генерации
    tasks = []
    old_employees = []
    
    for old_employee, equipment_list in grouped_equipment.items():
        old_employees.append(old_employee)
        task = generate_transfer_act_pdf(
            new_employee=new_employee,
            new_employee_dept=new_employee_dept,
            old_employee=old_employee,
            serials_data=equipment_list,
            db_name=db_name
        )
        tasks.append(task)
    
    # Выполняем параллельно с обработкой исключений
    results = await asyncio.gather(*tasks, return_exceptions=True)
    
    # Собираем информацию о результатах
    acts_info = []
    successful_count = 0
    failed_count = 0
    
    for idx, (old_employee, result) in enumerate(zip(old_employees, results)):
        equipment_list = grouped_equipment[old_employee]
        
        if isinstance(result, Exception):
            # Ошибка при генерации
            logger.error(f"Ошибка генерации акта для {old_employee}: {result}")
            acts_info.append({
                'old_employee': old_employee,
                'pdf_path': None,
                'filename': None,
                'equipment_count': len(equipment_list),
                'success': False,
                'error': str(result)
            })
            failed_count += 1
        elif result is None:
            # Генерация вернула None (ошибка)
            logger.error(f"Не удалось создать акт для {old_employee}")
            acts_info.append({
                'old_employee': old_employee,
                'pdf_path': None,
                'filename': None,
                'equipment_count': len(equipment_list),
                'success': False,
                'error': 'Ошибка генерации PDF'
            })
            failed_count += 1
        else:
            # Успешная генерация
            filename = os.path.basename(result)
            logger.info(f"Акт успешно создан для {old_employee}: {filename}")
            acts_info.append({
                'old_employee': old_employee,
                'pdf_path': result,
                'filename': filename,
                'equipment_count': len(equipment_list),
                'success': True,
                'error': None
            })
            successful_count += 1
    
    elapsed_time = time.time() - start_time
    logger.info(
        f"Генерация актов завершена: {successful_count}/{len(grouped_equipment)} успешно, "
        f"время: {elapsed_time:.2f}с"
    )
    
    if failed_count > 0:
        logger.warning(f"Не удалось создать {failed_count} актов")
    
    return acts_info
